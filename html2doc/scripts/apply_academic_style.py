#!/usr/bin/env python3
"""
Apply academic paper formatting to DOCX

Applies:
- Times New Roman font throughout
- Proper heading styles
- Centered images
- Full-width tables with borders
- Justified body text

Usage: python apply_academic_style.py [--input input.docx] [--output output.docx]
"""

import argparse
import os
import sys
import zipfile

try:
    from docx import Document
    from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
    from docx.shared import Pt, RGBColor, Inches
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
except ImportError:
    print("Installing python-docx...")
    os.system("pip install python-docx")
    from docx import Document
    from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
    from docx.shared import Pt, RGBColor, Inches
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn


# =============================================================================
# STYLING CONFIGURATION
# =============================================================================

STYLES = {
    # Fonts
    'font_family': 'Times New Roman',
    'body_font_size': 11,
    'title_font_size': 18,
    'author_font_size': 12,
    'affiliation_font_size': 10,
    'caption_font_size': 10,
    
    # Headings
    'heading1_size': 14,
    'heading2_size': 12,
    'heading3_size': 11,
    
    # Spacing
    'line_spacing': 1.5,
    'margin_inches': 1.0,
    
    # Table
    'table_header_bg': 'D9D9D9',  # Light gray
    'table_header_size': 9,
    'table_body_size': 10,
}


def set_table_border(table, color='000000', size=4):
    """Set table outer border."""
    tbl = table._element
    tblPr = tbl.xpath('.//w:tblPr')
    if not tblPr:
        tblPr = OxmlElement('w:tblPr')
        tbl.insert(0, tblPr)
    
    tblBorders = OxmlElement('w:tblBorders')
    for border_name in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
        border = OxmlElement(f'w:{border_name}')
        border.set(qn('w:val'), 'single')
        border.set(qn('w:sz'), str(size))
        border.set(qn('w:color'), color)
        tblBorders.append(border)
    
    tblPr.append(tblBorders)


def set_table_width(table, width_percent=100):
    """Set table to specified page width percentage."""
    tbl = table._element
    tblPr = tbl.xpath('.//w:tblPr')
    if not tblPr:
        tblPr = OxmlElement('w:tblPr')
        tbl.insert(0, tblPr)
    else:
        tblPr = tblPr[0]
    
    # Set table width as percentage
    tblW = OxmlElement('w:tblW')
    tblW.set(qn('w:type'), 'pct')
    tblW.set(qn('w:w'), str(width_percent * 50))  # 5000 = 100%
    tblPr.append(tblW)
    
    # Set cell spacing to 0
    tblCellSpacing = OxmlElement('w:tblCellSpacing')
    tblCellSpacing.set(qn('w:w'), '0')
    tblCellSpacing.set(qn('w:type'), 'dxa')
    tblPr.append(tblCellSpacing)


def format_table_cell(cell, is_header=False, is_first_col=False):
    """Format a table cell."""
    # Set shading for header
    if is_header:
        shading_elm = OxmlElement('w:shd')
        shading_elm.set(qn('w:fill'), STYLES['table_header_bg'])
        cell._tc.get_or_add_tcPr().append(shading_elm)
    
    # Format paragraphs
    font_size = STYLES['table_header_size'] if is_header else STYLES['table_body_size']
    
    for para in cell.paragraphs:
        for run in para.runs:
            run.font.name = STYLES['font_family']
            run.font.size = Pt(font_size)
            if is_header:
                run.font.bold = True
        
        if is_first_col:
            para.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
        else:
            para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        
        para.paragraph_format.space_after = Pt(2)


def has_drawing(element):
    """Check if paragraph contains a drawing (image)."""
    for elem in element.iter():
        if 'drawing' in elem.tag.lower():
            return True
    return False


def apply_academic_formatting(doc, table_width=100):
    """Apply academic formatting to document."""
    
    print("Applying academic formatting...")
    
    # Set document margins
    for section in doc.sections:
        section.top_margin = Inches(STYLES['margin_inches'])
        section.bottom_margin = Inches(STYLES['margin_inches'])
        section.left_margin = Inches(STYLES['margin_inches'])
        section.right_margin = Inches(STYLES['margin_inches'])
    
    # Process all paragraphs
    image_count = 0
    text_count = 0
    
    for para in doc.paragraphs:
        # Determine paragraph type
        is_heading = para.style.name.startswith('Heading') if para.style else False
        has_image = has_drawing(para._element)
        
        # Check for caption
        is_caption = 'Figure' in para.text or 'Table' in para.text
        
        if has_image:
            # Center images
            para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            image_count += 1
        
        elif is_caption:
            # Format captions
            para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            for run in para.runs:
                run.font.name = STYLES['font_family']
                run.font.size = Pt(STYLES['caption_font_size'])
                run.font.italic = True
        
        elif is_heading:
            # Format headings
            size_map = {
                'Heading 1': STYLES['heading1_size'],
                'Heading 2': STYLES['heading2_size'],
                'Heading 3': STYLES['heading3_size'],
            }
            size = size_map.get(para.style.name, STYLES['heading2_size'])
            
            for run in para.runs:
                run.font.name = STYLES['font_family']
                run.font.size = Pt(size)
                run.font.bold = True
                run.font.color.rgb = RGBColor(0, 0, 0)
            
            para.paragraph_format.space_before = Pt(12)
            para.paragraph_format.space_after = Pt(6)
        
        else:
            # Body text - justify
            if para.text.strip():
                para.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
                for run in para.runs:
                    run.font.name = STYLES['font_family']
                    run.font.size = Pt(STYLES['body_font_size'])
                text_count += 1
    
    print(f"  Centered {image_count} images")
    print(f"  Formatted {text_count} text paragraphs")
    
    # Apply table formatting
    print("  Formatting tables...")
    for idx, table in enumerate(doc.tables):
        print(f"    Table {idx + 1}: {len(table.rows)} rows x {len(table.columns)} cols")
        
        # Set borders
        set_table_border(table)
        
        # Set width
        set_table_width(table, table_width)
        
        # Format header row
        if table.rows:
            header_row = table.rows[0]
            for i, cell in enumerate(header_row.cells):
                format_table_cell(cell, is_header=True, is_first_col=(i == 0))
            
            # Format data rows
            for row in table.rows[1:]:
                for i, cell in enumerate(row.cells):
                    format_table_cell(cell, is_header=False, is_first_col=(i == 0))
    
    # Set normal style
    style = doc.styles['Normal']
    font = style.font
    font.name = STYLES['font_family']
    font.size = Pt(STYLES['body_font_size'])
    
    para_format = style.paragraph_format
    para_format.line_spacing = STYLES['line_spacing']
    para_format.space_after = Pt(0)
    
    print("  Formatting complete")


def main():
    parser = argparse.ArgumentParser(
        description='Apply academic paper formatting to DOCX'
    )
    parser.add_argument(
        '--input', '-i',
        default='paper_converted.docx',
        help='Input DOCX file (default: paper_converted.docx)'
    )
    parser.add_argument(
        '--output', '-o',
        default='paper_academic.docx',
        help='Output DOCX file (default: paper_academic.docx)'
    )
    parser.add_argument(
        '--table-width',
        type=int,
        default=100,
        help='Table width percentage (default: 100)'
    )
    
    args = parser.parse_args()
    
    if not os.path.exists(args.input):
        print(f"Error: Input file '{args.input}' not found")
        sys.exit(1)
    
    print(f"Loading: {args.input}")
    doc = Document(args.input)
    
    apply_academic_formatting(doc, args.table_width)
    
    doc.save(args.output)
    print(f"\nSaved: {args.output}")
    print("Done!")


if __name__ == '__main__':
    main()
